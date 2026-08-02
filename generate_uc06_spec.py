from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUTPUT = "UC-06_Dac-ta-use-case-Bao-cao.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    for index, line in enumerate(text.split("\n")):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    margins(cell)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"].font.size = Pt(12)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(12)
run = title.add_run("Bảng đặc tả use case UC-06 – Báo cáo")
run.bold = True
run.font.name = "Times New Roman"
run.font.size = Pt(15)

main_flow = """XEM DANH SÁCH BÁO CÁO
6.1.1.0. Sinh viên hoặc Quản trị viên mở chức năng Báo cáo từ menu hệ thống.
6.1.1.1. Frontend hiển thị giao diện danh sách báo cáo, bộ lọc loại báo cáo, trạng thái và thời gian gửi báo cáo.
6.1.1.2. Frontend gửi yêu cầu lấy danh sách báo cáo đến API.
6.1.1.3. Backend xác thực phiên đăng nhập, lấy định danh người dùng từ token hoặc session và kiểm tra quyền truy cập chức năng Báo cáo.
6.1.1.4. Backend xác định vai trò người dùng để giới hạn phạm vi dữ liệu được xem, ví dụ Sinh viên chỉ xem báo cáo của mình, Quản trị viên xem toàn bộ báo cáo.
6.1.1.5. Backend kiểm tra các tham số lọc như loại báo cáo, trạng thái xử lý, khoảng thời gian, từ khóa và phân trang.
6.1.1.6. Backend truy vấn bảng báo cáo, kết hợp dữ liệu người gửi, đối tượng bị báo cáo và trạng thái xử lý nếu cần.
6.1.1.7. Backend sắp xếp danh sách theo thời gian tạo mới nhất hoặc theo mức độ ưu tiên xử lý.
6.1.1.8. Backend chuyển đổi dữ liệu sang DTO, ẩn các thông tin nội bộ không cần hiển thị.
6.1.1.9. API trả về danh sách báo cáo cùng tổng số bản ghi và thông tin phân trang.
6.1.1.10. Frontend hiển thị danh sách báo cáo lên giao diện.
6.1.1.11. Người dùng có thể lọc, chuyển trang hoặc chọn một báo cáo để xem chi tiết.

XEM CHI TIẾT BÁO CÁO
6.2.1.0. Người dùng chọn một báo cáo trong danh sách.
6.2.1.1. Frontend lấy mã báo cáo được chọn.
6.2.1.2. Frontend gửi yêu cầu lấy chi tiết báo cáo đến API.
6.2.1.3. Backend xác thực người dùng và kiểm tra người dùng có quyền xem báo cáo đó hay không.
6.2.1.4. Backend truy vấn báo cáo theo mã và kiểm tra báo cáo còn tồn tại, chưa bị xóa hoặc không bị ẩn khỏi phạm vi truy cập.
6.2.1.5. Backend lấy thông tin người gửi báo cáo, loại đối tượng bị báo cáo, nội dung lý do, minh chứng kèm theo và trạng thái xử lý.
6.2.1.6. Backend truy vấn thêm dữ liệu đối tượng bị báo cáo, ví dụ hồ sơ người dùng, thông tin nhóm hoặc nội dung bài post.
6.2.1.7. Backend chuẩn hóa dữ liệu chi tiết, che thông tin nhạy cảm nếu người xem không có quyền quản trị.
6.2.1.8. API trả về thông tin chi tiết báo cáo.
6.2.1.9. Frontend hiển thị màn hình hoặc hộp thoại chi tiết báo cáo.
6.2.1.10. Người dùng xem được lý do báo cáo, loại báo cáo, đối tượng bị báo cáo, thời gian gửi và trạng thái xử lý.

BÁO CÁO NGƯỜI DÙNG
6.3.1.0. Sinh viên mở hồ sơ của một người dùng khác và chọn chức năng Báo cáo người dùng.
6.3.1.1. Frontend hiển thị form báo cáo gồm lý do, mô tả chi tiết và minh chứng nếu có.
6.3.1.2. Sinh viên nhập nội dung báo cáo và nhấn nút Gửi báo cáo.
6.3.1.3. Frontend kiểm tra dữ liệu bắt buộc, độ dài nội dung và định dạng tệp minh chứng nếu có.
6.3.1.4. Frontend gửi yêu cầu tạo báo cáo người dùng đến API.
6.3.1.5. Backend xác thực người gửi báo cáo và kiểm tra người dùng bị báo cáo có tồn tại hay không.
6.3.1.6. Backend kiểm tra người gửi không tự báo cáo chính mình và chưa gửi báo cáo trùng lặp trong khoảng thời gian giới hạn nếu hệ thống có quy định.
6.3.1.7. Backend kiểm tra lý do báo cáo thuộc danh mục hợp lệ và nội dung mô tả không rỗng.
6.3.1.8. Backend xử lý tệp minh chứng, kiểm tra dung lượng, định dạng và lưu đường dẫn minh chứng nếu hợp lệ.
6.3.1.9. Backend tạo bản ghi báo cáo với loại đối tượng là Người dùng, trạng thái ban đầu là Chờ xử lý.
6.3.1.10. Backend lưu báo cáo vào cơ sở dữ liệu, ghi nhận người gửi, người bị báo cáo, thời điểm gửi và dữ liệu minh chứng.
6.3.1.11. Backend có thể tạo thông báo hoặc log nghiệp vụ cho bộ phận quản trị xử lý báo cáo.
6.3.1.12. API trả về kết quả gửi báo cáo thành công.
6.3.1.13. Frontend hiển thị thông báo đã gửi báo cáo và đóng form.

BÁO CÁO NHÓM
6.4.1.0. Sinh viên mở trang nhóm học tập và chọn chức năng Báo cáo nhóm.
6.4.1.1. Frontend hiển thị form báo cáo nhóm gồm lý do, mô tả chi tiết và minh chứng nếu có.
6.4.1.2. Sinh viên nhập thông tin báo cáo và nhấn nút Gửi báo cáo.
6.4.1.3. Frontend kiểm tra dữ liệu bắt buộc, định dạng nội dung và tệp đính kèm.
6.4.1.4. Frontend gửi yêu cầu tạo báo cáo nhóm đến API.
6.4.1.5. Backend xác thực người gửi báo cáo và kiểm tra nhóm bị báo cáo có tồn tại hay không.
6.4.1.6. Backend kiểm tra trạng thái nhóm, ví dụ nhóm còn hoạt động, chưa bị xóa hoặc chưa bị khóa.
6.4.1.7. Backend kiểm tra người gửi có quyền truy cập nhóm hoặc có đủ điều kiện gửi báo cáo nhóm theo quy định hệ thống.
6.4.1.8. Backend kiểm tra lý do báo cáo, mô tả chi tiết và xử lý minh chứng kèm theo nếu có.
6.4.1.9. Backend tạo bản ghi báo cáo với loại đối tượng là Nhóm và liên kết đến mã nhóm bị báo cáo.
6.4.1.10. Backend lưu báo cáo vào cơ sở dữ liệu, ghi nhận trạng thái Chờ xử lý và thời điểm tạo.
6.4.1.11. Backend có thể cập nhật số lượng báo cáo của nhóm để hỗ trợ ưu tiên kiểm duyệt.
6.4.1.12. API trả về kết quả gửi báo cáo nhóm thành công.
6.4.1.13. Frontend hiển thị thông báo gửi báo cáo thành công.

BÁO CÁO BÀI POST
6.5.1.0. Sinh viên mở một bài post và chọn chức năng Báo cáo bài post.
6.5.1.1. Frontend hiển thị form báo cáo bài post gồm lý do, mô tả chi tiết và minh chứng nếu có.
6.5.1.2. Sinh viên nhập thông tin báo cáo và nhấn nút Gửi báo cáo.
6.5.1.3. Frontend kiểm tra dữ liệu bắt buộc, độ dài mô tả và định dạng tệp minh chứng nếu có.
6.5.1.4. Frontend gửi yêu cầu tạo báo cáo bài post đến API.
6.5.1.5. Backend xác thực người gửi báo cáo và kiểm tra bài post bị báo cáo có tồn tại hay không.
6.5.1.6. Backend kiểm tra bài post chưa bị xóa, chưa bị ẩn hoặc chưa bị xử lý trước đó theo chính sách hệ thống.
6.5.1.7. Backend kiểm tra quyền truy cập bài post, đặc biệt với bài post thuộc nhóm riêng tư hoặc phạm vi hạn chế.
6.5.1.8. Backend kiểm tra lý do báo cáo hợp lệ, mô tả chi tiết và xử lý minh chứng nếu có.
6.5.1.9. Backend tạo bản ghi báo cáo với loại đối tượng là Bài post và liên kết đến mã bài post bị báo cáo.
6.5.1.10. Backend lưu báo cáo vào cơ sở dữ liệu, ghi nhận người gửi, bài post bị báo cáo, trạng thái Chờ xử lý và thời điểm gửi.
6.5.1.11. Backend có thể tăng bộ đếm báo cáo của bài post hoặc gửi cảnh báo cho quản trị viên nếu số lượng báo cáo vượt ngưỡng.
6.5.1.12. API trả về kết quả gửi báo cáo bài post thành công.
6.5.1.13. Frontend hiển thị thông báo đã gửi báo cáo và giữ người dùng ở màn hình bài post hiện tại."""

alternative_flow = """6.1. Không tải được danh sách báo cáo (sau bước 6.1.1.2)
6.1.2. Backend hoặc cơ sở dữ liệu xảy ra lỗi trong quá trình truy vấn danh sách báo cáo.
6.1.3. API trả về trạng thái thất bại kèm thông báo lỗi phù hợp.
6.1.4. Frontend hiển thị thông báo không thể tải danh sách báo cáo.
6.1.5. Người dùng có thể tải lại dữ liệu hoặc thử lại sau.

6.2. Không tìm thấy báo cáo chi tiết (sau bước 6.2.1.2)
6.2.2. Backend không tìm thấy báo cáo theo mã yêu cầu hoặc báo cáo không thuộc phạm vi được phép xem.
6.2.3. API trả về lỗi không tồn tại hoặc không có quyền truy cập.
6.2.4. Frontend hiển thị thông báo không thể xem chi tiết báo cáo.
6.2.5. Người dùng quay lại danh sách báo cáo.

6.3. Dữ liệu báo cáo người dùng không hợp lệ (sau bước 6.3.1.4)
6.3.2. Backend phát hiện người dùng bị báo cáo không tồn tại, người gửi tự báo cáo chính mình hoặc nội dung báo cáo không hợp lệ.
6.3.3. API trả về lỗi xác thực dữ liệu.
6.3.4. Frontend hiển thị lỗi và yêu cầu Sinh viên chỉnh sửa thông tin báo cáo.
6.3.5. Sinh viên nhập lại thông tin và gửi lại báo cáo.

6.4. Dữ liệu báo cáo nhóm không hợp lệ (sau bước 6.4.1.4)
6.4.2. Backend phát hiện nhóm không tồn tại, đã bị xóa, bị khóa hoặc người gửi không đủ quyền báo cáo nhóm.
6.4.3. API trả về kết quả gửi báo cáo thất bại.
6.4.4. Frontend hiển thị thông báo không thể gửi báo cáo nhóm.
6.4.5. Sinh viên có thể kiểm tra lại nhóm hoặc thử lại sau.

6.5. Dữ liệu báo cáo bài post không hợp lệ (sau bước 6.5.1.4)
6.5.2. Backend phát hiện bài post không tồn tại, đã bị xóa, đã bị ẩn hoặc người gửi không có quyền truy cập bài post.
6.5.3. API trả về lỗi phù hợp với nguyên nhân thất bại.
6.5.4. Frontend hiển thị thông báo không thể gửi báo cáo bài post.
6.5.5. Sinh viên có thể quay lại danh sách bài post hoặc thử lại với bài post khác.

6.6. Tệp minh chứng không hợp lệ (sau các bước gửi báo cáo)
6.6.1. Backend phát hiện tệp minh chứng vượt dung lượng, sai định dạng hoặc không thể lưu trữ.
6.6.2. API từ chối yêu cầu tạo báo cáo và trả về thông báo lỗi.
6.6.3. Frontend hiển thị lỗi tại trường minh chứng.
6.6.4. Sinh viên thay đổi hoặc xóa tệp minh chứng rồi gửi lại."""

rows = [
    ("Thuộc tính", "Nội dung"),
    ("Tên use case", "Báo cáo"),
    ("Tác nhân chính", "Sinh viên, Quản trị viên"),
    ("Mô tả", "Cho phép người dùng xem danh sách báo cáo, xem chi tiết báo cáo và gửi báo cáo đối với người dùng, nhóm hoặc bài post vi phạm."),
    ("Tiền điều kiện", "Người dùng đã đăng nhập. Đối tượng bị báo cáo phải tồn tại trong hệ thống và người dùng có quyền truy cập đối tượng đó."),
    ("Hậu điều kiện", "Báo cáo được hiển thị hoặc được tạo mới thành công; dữ liệu báo cáo được lưu trong hệ thống với trạng thái xử lý phù hợp."),
    ("Luồng sự kiện chính", main_flow),
    ("Luồng thay thế/\nNgoại lệ", alternative_flow),
]

table = doc.add_table(rows=len(rows), cols=2)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

for index, (label, content) in enumerate(rows):
    left, right = table.rows[index].cells
    left.width = Cm(4.2)
    right.width = Cm(13.4)
    set_cell_text(left, label, bold=True)
    set_cell_text(right, content, bold=(index == 0))
    if index == 0:
        shade(left, "D9EAF7")
        shade(right, "D9EAF7")
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc.save(OUTPUT)
print(OUTPUT)
