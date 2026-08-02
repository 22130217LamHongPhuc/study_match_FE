from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUTPUT = "UC-04_Dac-ta-use-case-Quan-ly-lich-hoc_v4.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    for index, line in enumerate(text.split("\n")):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    margins(cell)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"].font.size = Pt(12)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(12)
run = title.add_run("Bảng đặc tả use case UC-04 – Quản lí lịch học")
run.bold = True
run.font.name = "Times New Roman"
run.font.size = Pt(15)

main_flow = """XEM LỊCH HỌC
4.1.1.0. Sinh viên mở chức năng Lịch học từ menu hoặc trang tổng quan.
4.1.1.1. Frontend hiển thị giao diện lịch với bộ lọc thời gian và khu vực danh sách lịch học.
4.1.1.2. Frontend gửi yêu cầu lấy dữ liệu lịch học của Sinh viên.
4.1.1.3. Backend xác thực phiên đăng nhập, lấy định danh Sinh viên từ token hoặc session và kiểm tra quyền truy cập.
4.1.1.4. Backend xác định phạm vi dữ liệu cần lấy, ví dụ theo ngày hiện tại, tuần hiện tại hoặc theo bộ lọc mà Frontend truyền lên.
4.1.1.5. Backend truy vấn bảng lịch học và các bảng liên quan nếu cần, ví dụ môn học, lớp học, phòng học hoặc trạng thái lịch.
4.1.1.6. Backend lọc các bản ghi theo chủ sở hữu, theo khoảng thời gian, theo trạng thái hiển thị và các điều kiện tìm kiếm hợp lệ.
4.1.1.7. Backend sắp xếp dữ liệu theo ngày, giờ bắt đầu, mức ưu tiên hoặc trạng thái để đảm bảo giao diện nhận dữ liệu đúng thứ tự.
4.1.1.8. Backend chuyển đổi dữ liệu sang DTO, loại bỏ trường nội bộ không cần hiển thị và chuẩn hóa định dạng ngày giờ.
4.1.1.9. API trả về danh sách lịch học cùng thông tin phân trang nếu có.
4.1.1.10. Frontend hiển thị các lịch học lên giao diện theo dạng danh sách hoặc lịch biểu.
4.1.1.11. Sinh viên có thể đổi chế độ xem theo ngày, tuần hoặc tháng để theo dõi lịch.
4.1.1.12. Use case tiếp tục chờ thao tác tiếp theo của Sinh viên.

TẠO LỊCH HỌC
4.2.1.0. Sinh viên nhấn nút Tạo lịch học.
4.2.1.1. Frontend mở form tạo mới với các trường thông tin cần nhập.
4.2.1.2. Sinh viên nhập tiêu đề, nội dung, thời gian bắt đầu, thời gian kết thúc, địa điểm và ghi chú.
4.2.1.3. Sinh viên chọn lưu lịch học.
4.2.1.4. Frontend kiểm tra dữ liệu bắt buộc, định dạng ngày giờ và tính hợp lệ của khoảng thời gian.
4.2.1.5. Frontend gửi yêu cầu POST tạo lịch học đến API.
4.2.1.6. Backend xác thực người dùng, kiểm tra quyền tạo lịch và chuẩn hóa dữ liệu đầu vào trước khi xử lý.
4.2.1.7. Backend kiểm tra các trường bắt buộc, kiểm tra thời gian bắt đầu nhỏ hơn thời gian kết thúc và xác nhận dữ liệu không rỗng.
4.2.1.8. Backend đối chiếu lịch mới với các lịch đã tồn tại để phát hiện trùng thời gian, trùng phòng học hoặc trùng nguồn dữ liệu nếu có liên kết.
4.2.1.9. Backend kiểm tra các khóa tham chiếu như lớp học, môn học, phòng học hoặc người tạo lịch có tồn tại hợp lệ hay không.
4.2.1.10. Nếu dữ liệu hợp lệ, backend tạo bản ghi lịch học mới trong một giao dịch để tránh ghi dở dang.
4.2.1.11. Backend lưu lịch học vào cơ sở dữ liệu, ghi nhận người tạo, thời điểm tạo và trạng thái ban đầu.
4.2.1.12. Backend có thể tạo log nghiệp vụ hoặc bản ghi audit để phục vụ tra cứu sau này.
4.2.1.13. API trả về kết quả tạo lịch học thành công cùng mã lịch học vừa sinh.
4.2.1.14. Frontend hiển thị thông báo thành công, đóng form và làm mới danh sách lịch học.
4.2.1.15. Lịch mới xuất hiện trong danh sách hoặc trên giao diện lịch.

XEM CHI TIẾT LỊCH HỌC
4.3.1.0. Sinh viên chọn một lịch học từ danh sách hoặc nhấn vào một ô lịch trên giao diện.
4.3.1.1. Frontend lấy mã lịch học được chọn.
4.3.1.2. Frontend gửi yêu cầu lấy chi tiết lịch học lên API.
4.3.1.3. Backend xác thực người dùng và kiểm tra quyền xem lịch học đó.
4.3.1.4. Backend truy vấn lịch học theo mã và kiểm tra bản ghi còn tồn tại hay đã bị xóa, ẩn hoặc vô hiệu hóa.
4.3.1.5. Backend lấy thêm thông tin liên quan như lớp học, môn học, phòng học, người tạo lịch và các nhãn trạng thái.
4.3.1.6. Backend chuẩn hóa dữ liệu chi tiết để trả về đúng định dạng hiển thị của Frontend.
4.3.1.7. Nếu lịch học không thuộc phạm vi được phép xem, backend trả về lỗi phân quyền.
4.3.1.8. API trả về dữ liệu chi tiết lịch học.
4.3.1.9. Frontend hiển thị màn hình hoặc hộp thoại chi tiết lịch học.
4.3.1.10. Sinh viên xem được đầy đủ nội dung, thời gian, địa điểm, ghi chú và trạng thái của lịch học.

TÌM KIẾM LỊCH HỌC
4.4.1.0. Sinh viên nhập từ khóa tìm kiếm hoặc chọn bộ lọc.
4.4.1.1. Frontend nhận giá trị tìm kiếm, chuẩn hóa nội dung và loại bỏ ký tự không cần thiết.
4.4.1.2. Sinh viên có thể tìm theo tiêu đề, ngày học, trạng thái hoặc địa điểm.
4.4.1.3. Frontend gửi yêu cầu tìm kiếm đến API.
4.4.1.4. Backend kiểm tra điều kiện lọc, chuẩn hóa từ khóa, chuyển đổi kiểu dữ liệu ngày giờ và xác nhận tham số phân trang nếu có.
4.4.1.5. Backend xây dựng câu truy vấn động theo các tiêu chí được truyền lên như từ khóa, khoảng ngày, trạng thái, lớp học hoặc địa điểm.
4.4.1.6. Backend áp dụng quyền truy cập để chỉ trả về các lịch học mà Sinh viên được phép xem.
4.4.1.7. Backend có thể giới hạn kết quả theo phân trang, sắp xếp theo độ liên quan hoặc theo thời gian gần nhất.
4.4.1.8. Backend chuyển dữ liệu kết quả sang danh sách DTO gọn nhẹ, phục vụ hiển thị và tìm chi tiết tiếp theo.
4.4.1.9. API trả về danh sách lịch học phù hợp.
4.4.1.10. Frontend hiển thị danh sách kết quả và cho phép Sinh viên chọn tiếp một lịch để xem chi tiết.
4.4.1.11. Sinh viên có thể xóa bộ lọc để quay về danh sách đầy đủ."""

alternative_flow = """4.1. Không tải được danh sách lịch học (sau bước 4.1.1.2)
4.1.2. Backend hoặc cơ sở dữ liệu trả về lỗi trong quá trình truy vấn, ví dụ lỗi kết nối, lỗi timeout hoặc lỗi xử lý logic.
4.1.3. API trả về trạng thái thất bại kèm thông tin lỗi phù hợp.
4.1.4. Frontend hiển thị thông báo không thể tải lịch học và có thể cung cấp nút thử lại.
4.1.5. Sinh viên có thể tải lại trang hoặc thử lại sau.

4.2. Dữ liệu tạo lịch học không hợp lệ (sau bước 4.2.1.4)
4.2.2. Frontend phát hiện thiếu trường bắt buộc hoặc sai định dạng thời gian.
4.2.3. Backend phát hiện lịch bị trùng, thời gian kết thúc nhỏ hơn thời gian bắt đầu, dữ liệu tham chiếu không tồn tại hoặc người dùng không có quyền tạo.
4.2.4. Backend trả về mã lỗi xác thực dữ liệu để Frontend hiển thị đúng thông báo.
4.2.5. Frontend hiển thị thông báo lỗi, đánh dấu các trường cần sửa và giữ lại dữ liệu hợp lệ đã nhập.
4.2.6. Sinh viên chỉnh sửa lại dữ liệu rồi gửi lại.

4.3. Không tìm thấy lịch học chi tiết (sau bước 4.3.1.2)
4.3.2. Backend không tìm thấy lịch học tương ứng với mã yêu cầu hoặc lịch đã bị xóa, ẩn, vô hiệu hóa.
4.3.3. API trả về lỗi không tồn tại hoặc không được phép truy cập.
4.3.4. Frontend hiển thị thông báo không tìm thấy lịch học.
4.3.5. Sinh viên quay lại màn hình danh sách.

4.4. Không có kết quả tìm kiếm phù hợp (sau bước 4.4.1.3)
4.4.2. Backend không tìm thấy lịch học nào khớp điều kiện tìm kiếm.
4.4.3. API trả về danh sách rỗng cùng tổng số kết quả bằng 0.
4.4.4. Frontend hiển thị trạng thái không có kết quả.
4.4.5. Sinh viên có thể thử lại với bộ lọc khác hoặc xóa bộ lọc để xem toàn bộ lịch."""

rows = [
    ("Thuộc tính", "Nội dung"),
    ("Tên use case", "Quản lí lịch học"),
    ("Tác nhân chính", "Sinh viên"),
    ("Mô tả", "Cho phép Sinh viên xem, tạo, xem chi tiết và tìm kiếm lịch học trong hệ thống."),
    ("Tiền điều kiện", "Sinh viên đã đăng nhập và có mã người dùng hợp lệ."),
    ("Hậu điều kiện", "Lịch học được hiển thị hoặc được tạo mới thành công; nếu có thay đổi thì dữ liệu được lưu vào hệ thống."),
    ("Luồng sự kiện chính", main_flow),
    ("Luồng thay thế/\nNgoại lệ", alternative_flow),
]

table = doc.add_table(rows=len(rows), cols=2)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

for index, (label, content) in enumerate(rows):
    left, right = table.rows[index].cells
    left.width = Cm(4.2)
    right.width = Cm(13.4)
    set_cell_text(left, label, bold=True)
    set_cell_text(right, content, bold=(index == 0))
    if index == 0:
        shade(left, "D9EAF7")
        shade(right, "D9EAF7")
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc.save(OUTPUT)
print(OUTPUT)
