from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT = Path(r"D:\Study_FE\frontend_studymatch\UC-03_Dac-ta-use-case-Dang-ky_v6.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_text(cell, text, bold=False, centered=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)


def add_lines(cell, lines):
    cell.text = ""
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(line)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_exception_content(cell):
    groups = [
        (
            "3.2. Dữ liệu đăng ký không hợp lệ (sau bước 3.1.1)",
            [
                "3.2.2. Frontend xác định thông tin đăng ký không hợp lệ.",
                "3.2.3. Frontend hiển thị thông báo lỗi và yêu cầu sinh viên nhập lại.",
                "3.2.4. Luồng quay lại bước 3.1.0.",
            ],
        ),
        (
            "3.3. Email đã được sử dụng (sau bước 3.1.3)",
            [
                "3.3.4. Backend xác định email đã tồn tại.",
                "3.3.5. API trả về kết quả đăng ký thất bại.",
                "3.3.6. Frontend hiển thị thông báo “Email đã được sử dụng”.",
                "3.3.7. Luồng quay lại bước 3.1.0.",
            ],
        ),
        (
            "3.4. Không thể tạo tài khoản (sau bước 3.1.4)",
            [
                "3.4.5. Backend không thể lưu tài khoản hoặc xử lý yêu cầu đăng ký.",
                "3.4.6. API trả về kết quả đăng ký thất bại.",
                "3.4.7. Frontend hiển thị thông báo đăng ký thất bại.",
                "3.4.8. Sinh viên có thể thực hiện đăng ký lại từ bước 3.1.0.",
            ],
        ),
        (
            "3.5. Mã xác thực không hợp lệ (sau bước 3.1.10)",
            [
                "3.5.11. Backend xác định mã không tồn tại, đã được sử dụng hoặc đã hết hạn.",
                "3.5.12. API từ chối xác thực và trả về thông báo xác thực không thành công.",
                "3.5.13. Sinh viên có thể yêu cầu gửi lại email xác thực.",
            ],
        ),
    ]

    cell.text = ""
    first = True
    for heading, steps in groups:
        paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        paragraph.paragraph_format.space_before = Pt(0 if paragraph is cell.paragraphs[0] else 5)
        paragraph.paragraph_format.space_after = Pt(0)
        heading_run = paragraph.add_run(heading)
        heading_run.bold = True
        heading_run.font.name = "Times New Roman"
        heading_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        heading_run.font.size = Pt(12)
        for step in steps:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(step)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.font.size = Pt(13)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(12)
title_run = title.add_run("Bảng đặc tả use case UC-03 – Đăng ký")
title_run.italic = True
title_run.font.name = "Times New Roman"
title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
title_run.font.size = Pt(14)

table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
table.columns[0].width = Cm(4.2)
table.columns[1].width = Cm(12.3)

header = table.rows[0]
header.cells[0].width = Cm(4.2)
header.cells[1].width = Cm(12.3)
set_cell_text(header.cells[0], "Thuộc tính", bold=True, centered=True)
set_cell_text(header.cells[1], "Nội dung", bold=True, centered=True)
set_cell_shading(header.cells[0], "E7E6E6")
set_cell_shading(header.cells[1], "E7E6E6")
for cell in header.cells:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

simple_rows = [
    ("Tên use case", "Đăng ký tài khoản"),
    ("Tác nhân chính", "Sinh viên"),
    (
        "Mô tả",
        "Cho phép sinh viên tạo tài khoản bằng họ tên, email và mật khẩu. Sau khi đăng ký, hệ thống gửi email chứa liên kết xác thực.",
    ),
    ("Tiền điều kiện", "Sinh viên chưa đăng nhập và email chưa được sử dụng trong hệ thống."),
    (
        "Hậu điều kiện",
        "Tài khoản sinh viên được tạo với trạng thái chưa xác thực email. Sau khi xác thực thành công, hệ thống cập nhật emailVerified = true.",
    ),
]

for label, content in simple_rows:
    row = table.add_row()
    keep_row_together(row)
    row.cells[0].width = Cm(4.2)
    row.cells[1].width = Cm(12.3)
    set_cell_text(row.cells[0], label, centered=True)
    set_cell_text(row.cells[1], content)
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

main_steps = [
    "3.1.0. Sinh viên nhập họ tên, email, mật khẩu, xác nhận mật khẩu và nhấn nút Tạo tài khoản.",
    "3.1.1. Frontend kiểm tra biểu mẫu đăng ký.",
    "3.1.2. Frontend gọi API đăng ký POST /api/auth/register.",
    "3.1.3. Backend kiểm tra email đã được sử dụng hay chưa.",
    "3.1.4. Backend mã hóa mật khẩu, tạo tài khoản sinh viên và lưu vào cơ sở dữ liệu.",
    "3.1.5. Backend tạo mã xác thực email có thời hạn 15 phút.",
    "3.1.6. Backend lưu mã xác thực và gửi email chứa liên kết xác thực đến sinh viên.",
    "3.1.7. API trả về kết quả đăng ký thành công cho Frontend.",
    "3.1.8. Frontend thông báo đăng ký thành công và chuyển sinh viên đến trang Đăng nhập.",
    "3.1.9. Sinh viên mở email và nhấn vào liên kết xác thực.",
    "3.1.10. API xác thực email GET /api/verify-email/confirm tiếp nhận yêu cầu; Backend kiểm tra mã xác thực.",
    "3.1.11. Backend xác thực email và đánh dấu mã xác thực đã được sử dụng.",
    "3.1.12. API trả về thông báo xác thực email thành công và cho phép sinh viên quay lại trang Đăng nhập.",
]

row = table.add_row()
row.cells[0].width = Cm(4.2)
row.cells[1].width = Cm(12.3)
set_cell_text(row.cells[0], "Luồng sự kiện chính", centered=True)
add_lines(row.cells[1], main_steps)
for cell in row.cells:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

row = table.add_row()
row.cells[0].width = Cm(4.2)
row.cells[1].width = Cm(12.3)
set_cell_text(row.cells[0], "Luồng thay thế/\nNgoại lệ", centered=True)
add_exception_content(row.cells[1])
for cell in row.cells:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

doc.save(OUTPUT)
print(OUTPUT)
