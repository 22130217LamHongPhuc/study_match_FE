from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT = "UC-01_Dac-ta-use-case-Quan-ly-ho-so-hoc-tap_v2.docx"

def setup_cell(cell, text, bold=False):
    cell.text = ""
    for i, line in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(line); r.bold = bold or line.startswith(("XEM HỒ SƠ HỌC TẬP", "CHỈNH SỬA HỒ SƠ HỌC TẬP")); r.font.name = "Times New Roman"; r.font.size = Pt(12)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    tcpr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "D9EAF7" if bold else "FFFFFF"); tcpr.append(shd)

main = """XEM HỒ SƠ HỌC TẬP
1.1.1.0. Sinh viên mở trang Hồ sơ học tập.
1.1.1.1. Frontend gọi GET /api/onboarding/profile/{userId}.
1.1.1.2. Backend tải thông tin hồ sơ, học kỳ, môn học, thời gian rảnh và lịch học từ cơ sở dữ liệu.
1.1.1.3. API trả về thông tin hồ sơ học tập đầy đủ.
1.1.1.4. Frontend hiển thị thông tin cá nhân, khóa/ngành, mục tiêu, GPA, tín chỉ, môn học, thời gian rảnh và lịch học.

CHỈNH SỬA HỒ SƠ HỌC TẬP
1.2.1.0. Sinh viên nhấn Chỉnh sửa hồ sơ.
1.2.1.1. Frontend mở biểu mẫu và điền sẵn dữ liệu hiện tại.
1.2.1.2. Sinh viên cập nhật thông tin và nhấn Lưu thay đổi.
1.2.1.3. Frontend kiểm tra dữ liệu và gọi PUT /api/profile/update.
1.2.1.4. Backend kiểm tra dữ liệu, cập nhật hồ sơ cá nhân/học tập và thay thế danh sách môn học, thời gian rảnh, lịch học.
1.2.1.5. API trả về hồ sơ đã cập nhật.
1.2.1.6. Frontend thông báo cập nhật thành công, tải lại dữ liệu và đóng biểu mẫu.
1.2.1.7. Use case kết thúc."""

alt = """1.2. Không tải được hồ sơ (sau bước 1.1.1)
1.2.2. Backend hoặc cơ sở dữ liệu trả về lỗi.
1.2.3. Frontend hiển thị thông báo không thể tải hồ sơ.
1.2.4. Sinh viên có thể tải lại trang.

1.3. Dữ liệu chỉnh sửa không hợp lệ (sau bước 1.1.8)
1.3.2. Backend từ chối dữ liệu bắt buộc hoặc dữ liệu không hợp lệ.
1.3.3. API trả về lỗi xác thực dữ liệu.
1.3.4. Frontend hiển thị lỗi và yêu cầu Sinh viên chỉnh sửa lại.
1.3.5. Luồng quay lại bước 1.1.7.

1.4. Không thể lưu hồ sơ (sau bước 1.1.9)
1.4.2. Backend không thể cập nhật cơ sở dữ liệu.
1.4.3. API trả về kết quả cập nhật thất bại.
1.4.4. Frontend hiển thị thông báo cập nhật thất bại; dữ liệu cũ vẫn được giữ nguyên.
1.4.5. Sinh viên có thể thử lưu lại từ bước 1.1.7."""

rows = [("Thuộc tính", "Nội dung"), ("Tên use case", "Quản lí hồ sơ học tập"), ("Tác nhân chính", "Sinh viên"),
("Mô tả", "Cho phép Sinh viên xem và chỉnh sửa thông tin hồ sơ học tập, gồm thông tin cá nhân, thông tin học tập, môn học, thời gian rảnh và lịch học."),
("Tiền điều kiện", "Sinh viên đã đăng nhập và có mã người dùng hợp lệ."),
("Hậu điều kiện", "Hồ sơ được hiển thị; nếu cập nhật thành công, thông tin mới được lưu vào hệ thống."),
("Luồng sự kiện chính", main), ("Luồng thay thế/\nNgoại lệ", alt)]

doc = Document(); sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Cm(1.8); sec.left_margin = sec.right_margin = Cm(1.8)
doc.styles["Normal"].font.name = "Times New Roman"; doc.styles["Normal"].font.size = Pt(12)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Bảng đặc tả use case UC-01 – Quản lí hồ sơ học tập"); r.bold=True; r.font.name="Times New Roman"; r.font.size=Pt(15)
t=doc.add_table(rows=len(rows), cols=2); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
for i,(a,b) in enumerate(rows):
    t.rows[i].cells[0].width=Cm(4.2); t.rows[i].cells[1].width=Cm(13.4)
    setup_cell(t.rows[i].cells[0],a, i==0); setup_cell(t.rows[i].cells[1],b, i==0)
doc.save(OUT); print(OUT)
