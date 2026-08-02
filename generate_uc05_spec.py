from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT = "UC-05_Dac-ta-use-case-Dang-nhap.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    for index, line in enumerate(text.split("\n")):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    margins(cell)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"].font.size = Pt(12)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(12)
run = title.add_run("Bảng đặc tả use case UC-05 – Đăng nhập")
run.bold = True
run.font.name = "Times New Roman"
run.font.size = Pt(15)

main_flow = """5.1.0. Người dùng nhập email, mật khẩu và nhấn nút Đăng nhập.
5.1.1. Frontend kiểm tra biểu mẫu đăng nhập.
5.1.2. Frontend gọi API đăng nhập phù hợp: POST /api/auth/login đối với Sinh viên hoặc POST /api/auth/admin/login đối với Admin.
5.1.3. Backend tìm kiếm tài khoản theo email.
5.1.4. Backend kiểm tra mật khẩu, trạng thái tài khoản và quyền đăng nhập của người dùng.
5.1.5. Backend cập nhật thời gian đăng nhập gần nhất của tài khoản.
5.1.6. Backend tạo access token và refresh token cho người dùng.
5.1.7. API trả về kết quả đăng nhập thành công cùng access token, refresh token, mã người dùng, trạng thái xác thực email, trạng thái onboarding và thông tin vai trò.
5.1.8. Frontend lưu access token, refresh token và mã người dùng vào bộ nhớ cục bộ.
5.1.9. Frontend chuyển người dùng đến trang phù hợp: Sinh viên đến trang xác thực email, onboarding hoặc trang chủ; Admin đến trang quản trị.
5.1.10. Use case kết thúc."""

alternative_flow = """5.2. Dữ liệu đăng nhập không hợp lệ (sau bước 5.1.1)
5.2.2. Frontend xác định email hoặc mật khẩu chưa được nhập hay không đúng định dạng.
5.2.3. Frontend hiển thị thông báo lỗi và yêu cầu người dùng nhập lại.
5.2.4. Luồng quay lại bước 5.1.0.

5.3. Không tìm thấy tài khoản (sau bước 5.1.3)
5.3.4. Backend xác định không tồn tại tài khoản tương ứng với email.
5.3.5. API trả về kết quả đăng nhập thất bại với mã lỗi USER_NOT_FOUND.
5.3.6. Frontend hiển thị thông báo không tìm thấy tài khoản.
5.3.7. Luồng quay lại bước 5.1.0.

5.4. Mật khẩu không chính xác (sau bước 5.1.4)
5.4.5. Backend xác định mật khẩu không khớp với mật khẩu của tài khoản.
5.4.6. API trả về kết quả đăng nhập thất bại với mã lỗi PASSWORD_INCORRECT.
5.4.7. Frontend hiển thị thông báo mật khẩu không chính xác.
5.4.8. Luồng quay lại bước 5.1.0.

5.5. Tài khoản bị khóa hoặc ngừng hoạt động (sau bước 5.1.4)
5.5.5. Backend xác định tài khoản đang bị khóa, đã bị xóa hoặc đã ngừng hoạt động.
5.5.6. API từ chối đăng nhập và trả về mã lỗi USER_LOCKED.
5.5.7. Frontend hiển thị thông báo tài khoản đã bị khóa hoặc ngừng hoạt động.
5.5.8. Use case kết thúc.

5.6. Sinh viên chưa xác thực email (sau bước 5.1.7)
5.6.8. Frontend xác định emailVerified = false.
5.6.9. Frontend chuyển Sinh viên đến trang xác thực email.
5.6.10. Sinh viên có thể thực hiện xác thực hoặc yêu cầu gửi lại email xác thực.

5.7. Sinh viên chưa hoàn thành onboarding (sau bước 5.1.7)
5.7.8. Frontend xác định email đã được xác thực nhưng onboardingCompleted = false.
5.7.9. Frontend chuyển Sinh viên đến trang onboarding.
5.7.10. Use case đăng nhập kết thúc; Sinh viên tiếp tục thực hiện quy trình onboarding.

5.8. Không thể xử lý yêu cầu đăng nhập (sau bước 5.1.6)
5.8.7. Backend không thể tạo token hoặc xử lý yêu cầu do lỗi hệ thống.
5.8.8. API trả về kết quả đăng nhập thất bại.
5.8.9. Frontend hiển thị thông báo đăng nhập thất bại và yêu cầu người dùng thử lại.
5.8.10. Người dùng có thể thực hiện đăng nhập lại từ bước 5.1.0."""

rows = [
    ("Thuộc tính", "Nội dung"),
    ("Tên use case", "Đăng nhập"),
    ("Tác nhân chính", "Sinh viên, Admin"),
    ("Mô tả", "Cho phép Sinh viên hoặc Admin đăng nhập vào StudyMatch bằng email và mật khẩu. Sau khi xác thực thành công, hệ thống cấp access token, refresh token và chuyển người dùng đến trang phù hợp với vai trò, trạng thái xác thực email và trạng thái onboarding."),
    ("Tiền điều kiện", "Người dùng chưa đăng nhập và đã có tài khoản trong hệ thống. Tài khoản không bị khóa, xóa hoặc ngừng hoạt động."),
    ("Hậu điều kiện", "Người dùng được xác thực thành công; hệ thống cấp access token và refresh token, cập nhật thời gian đăng nhập gần nhất. Frontend lưu thông tin phiên đăng nhập và chuyển người dùng đến trang phù hợp."),
    ("Luồng sự kiện chính", main_flow),
    ("Luồng thay thế/\nNgoại lệ", alternative_flow),
]

table = doc.add_table(rows=len(rows), cols=2)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

for index, (label, content) in enumerate(rows):
    left, right = table.rows[index].cells
    left.width = Cm(4.2)
    right.width = Cm(13.4)
    set_cell_text(left, label, bold=True)
    set_cell_text(right, content, bold=index == 0)
    if index == 0:
        shade(left, "D9EAF7")
        shade(right, "D9EAF7")
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc.save(OUTPUT)
print(OUTPUT)
