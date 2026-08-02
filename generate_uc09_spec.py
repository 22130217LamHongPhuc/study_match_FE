from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUTPUT = "UC-09_Dac-ta-use-case-Thong-ke-hoc-tap_v2.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    for index, line in enumerate(text.split("\n")):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    margins(cell)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"].font.size = Pt(12)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(12)
run = title.add_run("Bảng đặc tả use case UC-09 – Thống kê học tập")
run.bold = True
run.font.name = "Times New Roman"
run.font.size = Pt(15)

main_flow = """XEM THỐNG KÊ HOẠT ĐỘNG HỌC TẬP
9.1.1.0. Sinh viên chọn mục Thống kê học tập trên thanh menu chính.
9.1.1.1. Frontend mở trang Thống kê hoạt động học tập, hiển thị phần tiêu đề mô tả chức năng phân tích thời gian tích lũy, tỷ lệ chuyên cần và cơ cấu môn học/lịch học.
9.1.1.2. Frontend lấy thông tin người dùng hiện tại từ trạng thái đăng nhập hoặc local storage.
9.1.1.3. Frontend gửi yêu cầu lấy danh sách buổi học của Sinh viên đến API, sử dụng endpoint GET /api/study-sessions/user/{userId} với tham số phân trang và phạm vi thời gian nếu có.
9.1.1.4. Backend xác thực token, kiểm tra phiên đăng nhập và lấy userId của Sinh viên đang đăng nhập.
9.1.1.5. Backend kiểm tra userId trong đường dẫn có khớp với người dùng hiện tại hoặc có nằm trong phạm vi được phép truy cập hay không.
9.1.1.6. Backend kiểm tra các tham số lọc như sessionType, participantStatus, sessionStatus, startFrom, startTo, page và size.
9.1.1.7. Backend truy vấn dữ liệu buổi học của Sinh viên từ bảng study session và bảng participant/attendance liên quan.
9.1.1.8. Backend lấy các trường phục vụ thống kê gồm thời gian bắt đầu, thời gian kết thúc, loại buổi học, trạng thái buổi học, trạng thái tham gia và tổng thời lượng tham gia nếu có.
9.1.1.9. Backend chỉ trả về các buổi học thuộc Sinh viên, loại bỏ dữ liệu của người dùng khác và các bản ghi không còn hợp lệ.
9.1.1.10. Backend sắp xếp dữ liệu theo thời gian bắt đầu và trả về danh sách buổi học dạng PageResponse.
9.1.1.11. Frontend nhận dữ liệu, tính tổng số buổi học được mời/tham gia dựa trên danh sách buổi học trả về.
9.1.1.12. Frontend tính thời gian học tích lũy bằng cách cộng thời lượng tham gia của các buổi học đã ghi nhận attendance; nếu chưa có dữ liệu thì hiển thị 0 phút.
9.1.1.13. Frontend tính tỷ lệ chuyên cần bằng công thức số buổi Đã tham gia chia cho tổng số buổi học hợp lệ, không tính các buổi đã từ chối.
9.1.1.14. Frontend hiển thị ba thẻ chỉ số: Thời gian học tích lũy, Tỷ lệ chuyên cần và Tổng số buổi học.
9.1.1.15. Frontend hiển thị biểu đồ Cơ cấu Trạng thái Tham gia gồm các nhóm như Đã tham gia, Vắng mặt hoặc trạng thái khác theo dữ liệu thực tế.
9.1.1.16. Frontend hiển thị khung Đánh giá Chuyên cần với các mức Chăm chỉ, Trung bình và Báo động để Sinh viên hiểu ý nghĩa tỷ lệ chuyên cần.
9.1.1.17. Frontend hiển thị biểu đồ Thời gian học tập 30 ngày qua theo đơn vị phút.
9.1.1.18. Sinh viên xem được tổng quan hoạt động học tập của bản thân trên cùng một màn hình.

XEM CƠ CẤU TRẠNG THÁI THAM GIA
9.2.1.0. Sinh viên quan sát khối Cơ cấu Trạng thái Tham gia trên trang thống kê.
9.2.1.1. Frontend lấy danh sách buổi học đã tải và nhóm theo trạng thái tham gia.
9.2.1.2. Frontend xác định các trạng thái như COMPLETED, PARTIAL, ABSENT, NOT_JOINED hoặc các trạng thái tương đương từ dữ liệu backend.
9.2.1.3. Frontend chuyển trạng thái kỹ thuật sang nhãn tiếng Việt như Tham gia, Vắng mặt hoặc Tham gia một phần.
9.2.1.4. Frontend tính số lượng buổi học của từng trạng thái.
9.2.1.5. Frontend dựng biểu đồ tròn/donut và danh sách chú thích theo từng trạng thái.
9.2.1.6. Sinh viên xem được mình đã tham gia hoặc vắng mặt bao nhiêu buổi.

XEM ĐÁNH GIÁ CHUYÊN CẦN
9.3.1.0. Sinh viên xem khối Đánh giá Chuyên cần trên trang thống kê.
9.3.1.1. Frontend sử dụng tỷ lệ chuyên cần đã tính từ dữ liệu buổi học.
9.3.1.2. Frontend so sánh tỷ lệ chuyên cần với các ngưỡng hiển thị trên giao diện.
9.3.1.3. Nếu tỷ lệ từ 80% trở lên, Frontend hiển thị mức Chăm chỉ với ý nghĩa học tập chuyên cần, tham gia đầy đủ.
9.3.1.4. Nếu tỷ lệ từ 50% đến 79%, Frontend hiển thị mức Trung bình với ý nghĩa tham gia chưa đều, cần cố gắng hơn.
9.3.1.5. Nếu tỷ lệ dưới 50%, Frontend hiển thị mức Báo động với ý nghĩa vắng mặt quá nhiều, cần chủ động liên hệ nhóm học.
9.3.1.6. Sinh viên nắm được mức độ chuyên cần của bản thân dựa trên dữ liệu tham gia thực tế.

XEM BIỂU ĐỒ THỜI GIAN HỌC TẬP 30 NGÀY QUA
9.4.1.0. Sinh viên cuộn đến phần Thời gian học tập 30 ngày qua.
9.4.1.1. Frontend xác định mốc 30 ngày gần nhất tính từ ngày hiện tại.
9.4.1.2. Frontend lọc các buổi học có thời gian bắt đầu nằm trong 30 ngày gần nhất.
9.4.1.3. Frontend gom nhóm dữ liệu theo từng ngày.
9.4.1.4. Frontend cộng tổng thời lượng học tập trong ngày theo phút dựa trên attendance duration hoặc tổng thời lượng tham gia mà backend trả về.
9.4.1.5. Frontend tạo dữ liệu biểu đồ gồm nhãn ngày và số phút học tương ứng.
9.4.1.6. Frontend hiển thị biểu đồ thời gian học tập để Sinh viên theo dõi xu hướng học tập gần đây.
9.4.1.7. Sinh viên xem được ngày nào có hoạt động học tập và tổng số phút học của từng ngày.

LÀM MỚI DỮ LIỆU THỐNG KÊ
9.5.1.0. Sinh viên tải lại trang hoặc quay lại trang Thống kê học tập sau khi có buổi học mới.
9.5.1.1. Frontend gọi lại API lấy danh sách buổi học của Sinh viên.
9.5.1.2. Backend xác thực người dùng và truy vấn dữ liệu buổi học mới nhất.
9.5.1.3. Backend trả về danh sách buổi học đã cập nhật, bao gồm các trạng thái tham gia mới nhất sau khi Sinh viên vào hoặc rời phòng học.
9.5.1.4. Frontend tính lại thời gian học tích lũy, tỷ lệ chuyên cần, tổng số buổi học, cơ cấu trạng thái và biểu đồ 30 ngày.
9.5.1.5. Frontend cập nhật lại toàn bộ số liệu trên màn hình.
9.5.1.6. Sinh viên xem được thống kê học tập mới nhất."""

alternative_flow = """9.1. Chưa đăng nhập hoặc phiên đăng nhập không hợp lệ (sau bước 9.1.1.3)
9.1.2. Backend xác định token không tồn tại, hết hạn hoặc không hợp lệ.
9.1.3. API trả về lỗi xác thực.
9.1.4. Frontend chuyển Sinh viên về trang Đăng nhập hoặc hiển thị thông báo yêu cầu đăng nhập lại.

9.2. Không có quyền xem dữ liệu thống kê (sau bước 9.1.1.5)
9.2.2. Backend phát hiện userId yêu cầu không thuộc người dùng hiện tại hoặc người dùng không có quyền truy cập dữ liệu này.
9.2.3. API trả về lỗi không có quyền truy cập.
9.2.4. Frontend hiển thị thông báo không thể xem thống kê học tập.
9.2.5. Sinh viên quay lại trang trước hoặc trang chủ.

9.3. Không tải được dữ liệu buổi học (sau bước 9.1.1.7)
9.3.2. Backend hoặc cơ sở dữ liệu xảy ra lỗi trong quá trình truy vấn danh sách buổi học.
9.3.3. API trả về trạng thái thất bại kèm thông báo lỗi phù hợp.
9.3.4. Frontend hiển thị trạng thái lỗi hoặc thông báo không thể tải thống kê học tập.
9.3.5. Sinh viên có thể tải lại trang để thử lại.

9.4. Không có dữ liệu học tập (sau bước 9.1.1.10)
9.4.2. Backend trả về danh sách buổi học rỗng hoặc không có bản ghi hợp lệ để thống kê.
9.4.3. Frontend hiển thị các chỉ số mặc định như 0 phút, 0%, 0 buổi học.
9.4.4. Frontend hiển thị biểu đồ trống hoặc thông báo chưa có dữ liệu học tập.
9.4.5. Sinh viên tiếp tục sử dụng chức năng lịch học/phòng học để phát sinh dữ liệu thống kê.

9.5. Dữ liệu thời lượng tham gia chưa được ghi nhận (sau bước 9.1.1.12)
9.5.2. Frontend không tìm thấy duration hoặc totalDurationSeconds của một số buổi học.
9.5.3. Frontend bỏ qua các bản ghi thiếu thời lượng hoặc tính giá trị mặc định bằng 0 phút.
9.5.4. Frontend vẫn hiển thị tổng số buổi học và tỷ lệ chuyên cần nếu có trạng thái tham gia.
9.5.5. Sinh viên có thể xem thống kê cập nhật sau khi hệ thống ghi nhận lại attendance.

9.6. Tham số lọc/phân trang không hợp lệ (sau bước 9.1.1.6)
9.6.2. Backend phát hiện tham số page, size, startFrom, startTo hoặc trạng thái lọc không hợp lệ.
9.6.3. API trả về lỗi xác thực dữ liệu.
9.6.4. Frontend hiển thị thông báo lỗi hoặc quay về bộ lọc mặc định.
9.6.5. Frontend gọi lại API với tham số hợp lệ."""

rows = [
    ("Thuộc tính", "Nội dung"),
    ("Tên use case", "Thống kê học tập"),
    ("Tác nhân chính", "Sinh viên"),
    ("Mô tả", "Cho phép Sinh viên xem thống kê hoạt động học tập của bản thân gồm thời gian học tích lũy, tỷ lệ chuyên cần, tổng số buổi học, cơ cấu trạng thái tham gia và biểu đồ thời gian học tập 30 ngày qua."),
    ("Tiền điều kiện", "Sinh viên đã đăng nhập. Hệ thống có dữ liệu lịch học/buổi học và dữ liệu trạng thái tham gia hoặc attendance để tính thống kê."),
    ("Hậu điều kiện", "Các chỉ số và biểu đồ thống kê học tập được hiển thị theo dữ liệu mới nhất; nếu không có dữ liệu thì hệ thống hiển thị giá trị mặc định hoặc trạng thái trống."),
    ("Luồng sự kiện chính", main_flow),
    ("Luồng thay thế/\nNgoại lệ", alternative_flow),
]

table = doc.add_table(rows=len(rows), cols=2)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

for index, (label, content) in enumerate(rows):
    left, right = table.rows[index].cells
    left.width = Cm(4.2)
    right.width = Cm(13.4)
    set_cell_text(left, label, bold=True)
    set_cell_text(right, content, bold=(index == 0))
    if index == 0:
        shade(left, "D9EAF7")
        shade(right, "D9EAF7")
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc.save(OUTPUT)
print(OUTPUT)
